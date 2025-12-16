# services/multimodal_analysis.py
from typing import List, Dict, Any, Optional, Tuple
import os
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
import hashlib
import json
from pathlib import Path
import re
import pandas as pd # Added for spreadsheet analysis

@dataclass
class TableData:
    """Structured table data"""
    headers: List[str]
    rows: List[List[str]]

logger = logging.getLogger(__name__)

@dataclass
class ForensicResult:
    """Forensic analysis result"""
    evidence_id: str
    file_type: str
    manipulation_score: float  # 0-100, higher = more likely manipulated
    authenticity_score: float  # 0-100, higher = more authentic
    forensic_indicators: List[str]
    metadata_analysis: Dict[str, Any]
    confidence: float
    analysis_timestamp: datetime

@dataclass
class MultiModalAnalysis:
    """Complete multi-modal analysis result"""
    evidence_id: str
    file_path: str
    file_type: str
    size_bytes: int
    
    # Text analysis
    extracted_text: str
    key_entities: List[Dict[str, Any]]
    sentiment_score: float
    language_detected: str
    
    # Visual analysis (for images)
    visual_features: Dict[str, Any]
    objects_detected: List[str]
    faces_detected: List[Dict[str, Any]]
    
    # Document analysis (for documents)
    document_structure: Dict[str, Any]
    signatures_detected: List[Dict[str, Any]]
    form_fields: List[Dict[str, Any]]
    extracted_tables: List[TableData]
    
    # Forensic analysis
    forensic_result: ForensicResult
    
    # Quality assessment
    quality_score: float
    relevance_score: float
    admissibility_score: float
    
    # Processing metadata
    processing_time: float
    analysis_timestamp: datetime
    errors: List[str]

    document_type: str = "general_document"
    bank_statement_data: Optional[Dict[str, Any]] = None
    expense_data: Optional[Dict[str, Any]] = None

class MultiModalAnalyzer:
    """Advanced multi-modal evidence analysis with forensic capabilities"""
    
    def __init__(self, temp_dir: str = None):
        self.temp_dir = temp_dir or os.path.join(os.getcwd(), 'temp_analysis')
        os.makedirs(self.temp_dir, exist_ok=True)
        
        # Initialize analysis modules
        self._init_modules()
    
    def _check_ocr_availability(self) -> bool:
        """Check if OCR is available"""
        try:
            import pytesseract
            import cv2
            return True
        except ImportError:
            logger.warning("OCR not available - install pytesseract and opencv-python")
            return False
    
    def _check_image_analysis_available(self) -> bool:
        """Check if image analysis is available"""
        try:
            from PIL import Image
            import cv2
            import numpy as np
            return True
        except ImportError:
            logger.warning("Image analysis not available - install Pillow and opencv-python")
            return False
    
    def _check_document_analysis_available(self) -> bool:
        """Check if document analysis is available"""
        try:
            import PyPDF2
            import docx
            import pandas as pd # Added for Excel/CSV analysis
            return True
        except ImportError:
            logger.warning("Document analysis not available - install PyPDF2, python-docx, and pandas")
            return False
    
    def _check_forensic_availability(self) -> bool:
        """Check if forensic analysis is available"""
        try:
            import cv2
            import numpy as np
            from PIL import Image
            return True
        except ImportError:
            logger.warning("Forensic analysis not available - install required dependencies")
            return False

    def _init_modules(self):
        """Initialize analysis modules"""
        self.ocr_available = self._check_ocr_availability()
        self.image_analysis_available = self._check_image_analysis_available()
        self.document_analysis_available = self._check_document_analysis_available()
        self.forensic_available = self._check_forensic_availability()

        logger.info(f"Analysis modules - OCR: {self.ocr_available}, "
                   f"Image: {self.image_analysis_available}, "
                   f"Document: {self.document_analysis_available}, "
                   f"Forensic: {self.forensic_available}")

    def analyze_evidence(self, file_path: str, options: Dict[str, Any] = None) -> MultiModalAnalysis:
        """
        Perform comprehensive multi-modal analysis on evidence file
        
        Args:
            file_path: Path to evidence file
            options: Analysis options
            
        Returns:
            Complete multi-modal analysis result
        """
        start_time = datetime.now(timezone.utc)
        options = options or {}
        
        try:
            # Basic file information
            evidence_id = self._generate_evidence_id(file_path)
            file_type = self._detect_file_type(file_path)
            file_size = os.path.getsize(file_path)
            
            # Initialize analysis result
            analysis = MultiModalAnalysis(
                evidence_id=evidence_id,
                file_path=file_path,
                file_type=file_type,
                size_bytes=file_size,
                extracted_text="",
                key_entities=[],
                sentiment_score=0.0,
                language_detected="unknown",
                visual_features={},
                objects_detected=[],
                faces_detected=[],
                document_structure={},
                signatures_detected=[],
                form_fields=[],
                extracted_tables=[],
                forensic_result=None,
                quality_score=0.0,
                relevance_score=0.0,
                admissibility_score=0.0,
                processing_time=0.0,
                analysis_timestamp=start_time,
                errors=[]
            )
            
            # Perform analysis based on file type
            if file_type.startswith('image/'):
                self._analyze_image_file(file_path, analysis, options)
            elif file_type == 'application/pdf':
                self._analyze_pdf_file(file_path, analysis, options)
            elif file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                self._analyze_word_file(file_path, analysis, options)
            elif file_type in ['text/csv', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
                self._analyze_spreadsheet_file(file_path, analysis)
            elif file_type.startswith('text/'):
                self._analyze_text_file(file_path, analysis, options)
            else:
                analysis.errors.append(f"Unsupported file type for analysis: {file_type}")
            
            # Perform forensic analysis if available
            if self.forensic_available and options.get('enable_forensics', True):
                analysis.forensic_result = self._perform_forensic_analysis(file_path, analysis)

            # Identify document type and map specific data
            analysis.document_type = self._identify_document_type(analysis)
            if analysis.document_type == "bank_statement":
                analysis.bank_statement_data = self._map_bank_statement_data(analysis)
            elif analysis.document_type == "expense_report":
                analysis.expense_data = self._map_expense_data(analysis)
            
            # Calculate quality scores
            self._calculate_quality_scores(analysis)
            
            # Calculate processing time
            analysis.processing_time = (datetime.now(timezone.utc) - start_time).total_seconds()
            
            logger.info(f"Multi-modal analysis completed for {evidence_id} in {analysis.processing_time:.2f}s")
            
            return analysis
            
        except Exception as e:
            logger.error(f"Multi-modal analysis failed for {file_path}: {str(e)}")
            raise
    
    def _analyze_image_file(self, file_path: str, analysis: MultiModalAnalysis, options: Dict[str, Any]):
        """Analyze image file with OCR and visual analysis"""
        if not self.image_analysis_available:
            analysis.errors.append("Image analysis not available")
            return
        
        try:
            from PIL import Image
            import cv2
            import numpy as np
            
            # Open image
            image = Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Basic visual features
            analysis.visual_features = {
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'mode': image.mode,
                'has_transparency': image.mode in ('RGBA', 'LA') or 'transparency' in image.info
            }
            
            # OCR text extraction
            if self.ocr_available and options.get('enable_ocr', True):
                analysis.extracted_text = self._extract_text_from_image(file_path)
                if analysis.extracted_text:
                    analysis.key_entities = self._extract_entities_from_text(analysis.extracted_text)
                    analysis.sentiment_score = self._analyze_sentiment(analysis.extracted_text)
                    analysis.language_detected = self._detect_language(analysis.extracted_text)
            
            # Object detection (simplified)
            if options.get('enable_object_detection', False):
                analysis.objects_detected = self._detect_objects_in_image(file_path)
            
            # Face detection (simplified)
            if options.get('enable_face_detection', False):
                analysis.faces_detected = self._detect_faces_in_image(file_path)
            
        except Exception as e:
            analysis.errors.append(f"Image analysis failed: {str(e)}")
    
    def _analyze_pdf_file(self, file_path: str, analysis: MultiModalAnalysis, options: Dict[str, Any]):
        """Analyze PDF file with text extraction and structure analysis"""
        if not self.document_analysis_available:
            analysis.errors.append("Document analysis not available")
            return
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text_content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                
                analysis.extracted_text = text_content.strip()
                
                if analysis.extracted_text:
                    analysis.key_entities = self._extract_entities_from_text(analysis.extracted_text)
                    analysis.sentiment_score = self._analyze_sentiment(analysis.extracted_text)
                    analysis.language_detected = self._detect_language(analysis.extracted_text)
                
                # Document structure analysis
                analysis.document_structure = {
                    'page_count': len(pdf_reader.pages),
                    'has_forms': any('/AcroForm' in page for page in pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                    'has_signatures': any('/V' in page.get('/Annots', {}) for page in pdf_reader.pages if hasattr(page, 'get'))
                }
                
                # Extract metadata
                if pdf_reader.metadata:
                    analysis.document_structure['metadata'] = dict(pdf_reader.metadata)

            # Basic attempt to extract tables from text content
            if analysis.extracted_text:
                tables = self._extract_tables_from_text(analysis.extracted_text)
                analysis.extracted_tables.extend(tables)
            
        except Exception as e:
            analysis.errors.append(f"PDF analysis failed: {str(e)}")
    
    def _analyze_word_file(self, file_path: str, analysis: MultiModalAnalysis, options: Dict[str, Any]):
        """Analyze Word document"""
        if not self.document_analysis_available:
            analysis.errors.append("Document analysis not available")
            return
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            analysis.extracted_text = text_content.strip()
            
            if analysis.extracted_text:
                analysis.key_entities = self._extract_entities_from_text(analysis.extracted_text)
                analysis.sentiment_score = self._analyze_sentiment(analysis.extracted_text)
                analysis.language_detected = self._detect_language(analysis.extracted_text)
            
            # Document structure
            analysis.document_structure = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'has_comments': len(doc.comments) > 0 if hasattr(doc, 'comments') else False,
                'has_track_changes': False  # Would need additional library
            }
            
            # Extract tables
            analysis.extracted_tables = self._extract_tables_from_word(doc)
            
        except Exception as e:
            analysis.errors.append(f"Word document analysis failed: {str(e)}")
    
    def _analyze_text_file(self, file_path: str, analysis: MultiModalAnalysis, options: Dict[str, Any]):
        """Analyze text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text_content = file.read()
            
            analysis.extracted_text = text_content.strip()
            
            if analysis.extracted_text:
                analysis.key_entities = self._extract_entities_from_text(analysis.extracted_text)
                analysis.sentiment_score = self._analyze_sentiment(analysis.extracted_text)
                analysis.language_detected = self._detect_language(analysis.extracted_text)
            
            analysis.document_structure = {
                'character_count': len(text_content),
                'word_count': len(text_content.split()),
                'line_count': len(text_content.splitlines())
            }
            
        except Exception as e:
            analysis.errors.append(f"Text file analysis failed: {str(e)}")
    
    def _extract_tables_from_text(self, text: str) -> List[TableData]:
        """Attempt to extract tables from raw text content using basic heuristics"""
        tables: List[TableData] = []
        lines = text.splitlines()
        current_table_lines: List[str] = []
        
        for line in lines:
            stripped_line = line.strip()
            
            # Heuristic: if a line looks like it could be part of a table (e.g., multiple columns, not too short)
            # This is a very basic heuristic and can be improved with more advanced parsing.
            if len(stripped_line.split()) > 2 and len(stripped_line) > 20:
                current_table_lines.append(stripped_line)
            else:
                if len(current_table_lines) > 1: # Require at least a header and one row
                    # Attempt to parse table from current_table_lines
                    headers, rows = self._parse_text_table(current_table_lines)
                    if headers or rows:
                        tables.append(TableData(headers=headers, rows=rows))
                current_table_lines = []
        
        # Process any remaining table lines at the end of the document
        if len(current_table_lines) > 1:
            headers, rows = self._parse_text_table(current_table_lines)
            if headers or rows:
                tables.append(TableData(headers=headers, rows=rows))
                
        return tables

    def _parse_text_table(self, table_lines: List[str]) -> Tuple[List[str], List[List[str]]]:
        """Parse a list of text lines into table headers and rows using whitespace/delimiter heuristics"""
        headers: List[str] = []
        rows: List[List[str]] = []

        if not table_lines: return headers, rows

        # Attempt to detect a consistent delimiter (e.g., multiple spaces, tabs, commas)
        # For simplicity, we'll try to split by multiple spaces or tabs first
        delimiter = '   ' # three spaces as a common column separator
        if '\t' in table_lines[0]:
            delimiter = '\t'
        elif ',' in table_lines[0] and len(table_lines[0].split(',')) > 2:
            delimiter = ','

        # Assume the first line is headers if it seems reasonable
        first_line_cells = [cell.strip() for cell in table_lines[0].split(delimiter) if cell.strip()]
        if all(word[0].isupper() or not word.isalpha() for word in first_line_cells if word): # Heuristic for headers
            headers = first_line_cells
            data_lines = table_lines[1:]
        else:
            data_lines = table_lines

        for line in data_lines:
            cells = [cell.strip() for cell in line.split(delimiter) if cell.strip()]
            if cells: rows.append(cells)

        # Basic post-processing to align columns if headers were not detected perfectly
        if headers and not rows: # If only headers were detected, but no rows matched the delimiter logic, re-evaluate
            pass # More advanced logic needed here for robust table parsing

        return headers, rows
    
    def _extract_tables_from_word(self, doc: 'docx.Document') -> List[TableData]:
        """Extract tables from a Word document"""
        extracted_tables: List[TableData] = []
        for table in doc.tables:
            headers: List[str] = []
            rows: List[List[str]] = []

            # Extract headers (first row)
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]

            # Extract data rows (all rows after the first)
            for i, row in enumerate(table.rows):
                if i == 0 and headers:  # Skip header row if already extracted
                    continue
                row_data = [cell.text.strip() for cell in row.cells]
                rows.append(row_data)
            
            if headers or rows: # Only add if table has content
                extracted_tables.append(TableData(headers=headers, rows=rows))
                
        return extracted_tables

    def _perform_forensic_analysis(self, file_path: str, analysis: MultiModalAnalysis) -> ForensicResult:
        """Perform forensic analysis on the file"""
        try:
            if analysis.file_type.startswith('image/'):
                return self._analyze_image_forensics(file_path, analysis)
            elif analysis.file_type == 'application/pdf':
                return self._analyze_pdf_forensics(file_path, analysis)
            else:
                return self._analyze_general_forensics(file_path, analysis)
                
        except Exception as e:
            logger.error(f"Forensic analysis failed: {str(e)}")
            return ForensicResult(
                evidence_id=analysis.evidence_id,
                file_type=analysis.file_type,
                manipulation_score=0.0,
                authenticity_score=50.0,
                forensic_indicators=[f"Forensic analysis failed: {str(e)}"],
                metadata_analysis={},
                confidence=0.0,
                analysis_timestamp=datetime.now(timezone.utc)
            )
    
    def _analyze_image_forensics(self, file_path: str, analysis: MultiModalAnalysis) -> ForensicResult:
        """Perform forensic analysis on image"""
        try:
            from PIL import Image
            import cv2
            import numpy as np
            
            image = Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            forensic_indicators = []
            manipulation_score = 0.0
            authenticity_score = 100.0
            
            # Error Level Analysis (ELA)
            ela_result = self._perform_ela(image)
            if ela_result['score'] > 15:
                forensic_indicators.append("High error level analysis score - possible manipulation")
                manipulation_score += 25
                authenticity_score -= 20
            
            # Noise analysis
            noise_result = self._analyze_noise_pattern(image)
            if noise_result['inconsistent']:
                forensic_indicators.append("Inconsistent noise pattern detected")
                manipulation_score += 20
                authenticity_score -= 15
            
            # Metadata analysis
            metadata_result = self._analyze_image_metadata(image)
            if metadata_result['suspicious']:
                forensic_indicators.extend(metadata_result['indicators'])
                manipulation_score += 15
                authenticity_score -= 10
            
            # Clone detection
            clone_result = self._detect_clone_regions(image)
            if clone_result['clones_detected']:
                forensic_indicators.append("Clone regions detected")
                manipulation_score += 30
                authenticity_score -= 25
            
            # Cap scores
            manipulation_score = min(manipulation_score, 100.0)
            authenticity_score = max(authenticity_score, 0.0)
            
            return ForensicResult(
                evidence_id=analysis.evidence_id,
                file_type=analysis.file_type,
                manipulation_score=manipulation_score,
                authenticity_score=authenticity_score,
                forensic_indicators=forensic_indicators,
                metadata_analysis={
                    'ela_score': ela_result['score'],
                    'noise_analysis': noise_result,
                    'metadata_analysis': metadata_result,
                    'clone_detection': clone_result
                },
                confidence=0.8,
                analysis_timestamp=datetime.now(timezone.utc)
            )
            
        except Exception as e:
            logger.error(f"Image forensic analysis failed: {str(e)}")
            raise
    
    def _analyze_pdf_forensics(self, file_path: str, analysis: MultiModalAnalysis) -> ForensicResult:
        """Perform forensic analysis on PDF"""
        try:
            import PyPDF2
            
            forensic_indicators = []
            manipulation_score = 0.0
            authenticity_score = 100.0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check for modifications
                if pdf_reader.is_encrypted:
                    forensic_indicators.append("PDF is encrypted")
                    manipulation_score += 10
                
                # Check for suspicious metadata
                if pdf_reader.metadata:
                    metadata = dict(pdf_reader.metadata)
                    if self._has_suspicious_pdf_metadata(metadata):
                        forensic_indicators.append("Suspicious PDF metadata detected")
                        manipulation_score += 15
                        authenticity_score -= 10
                
                # Check for form fields (potential for manipulation)
                for page in pdf_reader.pages:
                    if '/AcroForm' in page:
                        forensic_indicators.append("PDF contains form fields")
                        manipulation_score += 5
                
            return ForensicResult(
                evidence_id=analysis.evidence_id,
                file_type=analysis.file_type,
                manipulation_score=manipulation_score,
                authenticity_score=authenticity_score,
                forensic_indicators=forensic_indicators,
                metadata_analysis={'pdf_metadata': metadata if 'metadata' in locals() else {}},
                confidence=0.7,
                analysis_timestamp=datetime.now(timezone.utc)
            )
            
        except Exception as e:
            logger.error(f"PDF forensic analysis failed: {str(e)}")
            raise
    
    def _analyze_general_forensics(self, file_path: str, analysis: MultiModalAnalysis) -> ForensicResult:
        """Perform general forensic analysis for non-image/PDF files"""
        forensic_indicators = []
        manipulation_score = 0.0
        authenticity_score = 100.0
        
        # File metadata analysis
        try:
            import magic
            file_mime = magic.from_file(file_path, mime=True)
            declared_mime = analysis.file_type
            
            if file_mime != declared_mime:
                forensic_indicators.append(f"File type mismatch: declared {declared_mime}, actual {file_mime}")
                manipulation_score += 20
                authenticity_score -= 15
                
        except ImportError:
            logger.warning("python-magic not available for file type verification")
        
        return ForensicResult(
            evidence_id=analysis.evidence_id,
            file_type=analysis.file_type,
            manipulation_score=manipulation_score,
            authenticity_score=authenticity_score,
            forensic_indicators=forensic_indicators,
            metadata_analysis={},
            confidence=0.5,
            analysis_timestamp=datetime.now(timezone.utc)
        )
    
    def _perform_ela(self, image: 'Image') -> Dict[str, Any]:
        """Perform Error Level Analysis"""
        try:
            import cv2
            import numpy as np
            import tempfile
            
            # Save with high quality
            temp_path = os.path.join(self.temp_dir, f"temp_ela_{hash(image.tobytes())}.jpg")
            image.save(temp_path, 'JPEG', quality=95)
            
            # Reload and compare
            reloaded = Image.open(temp_path)
            original_array = np.array(image)
            reloaded_array = np.array(reloaded)
            
            # Calculate difference
            if original_array.shape == reloaded_array.shape:
                diff = np.abs(original_array.astype(np.int16) - reloaded_array.astype(np.int16))
                ela_score = np.mean(diff)
                
                # Clean up
                os.unlink(temp_path)
                
                return {'score': float(ela_score)}
            else:
                return {'score': 0.0}
                
        except Exception as e:
            logger.error(f"ELA analysis failed: {str(e)}")
            return {'score': 0.0}
    
    def _analyze_noise_pattern(self, image: 'Image') -> Dict[str, Any]:
        """Analyze noise patterns in image"""
        try:
            import cv2
            import numpy as np
            
            # Convert to numpy array
            img_array = np.array(image)
            
            # Calculate noise in different regions
            h, w = img_array.shape[:2]
            regions = [
                img_array[0:h//4, 0:w//4],
                img_array[0:h//4, w//4:w//2],
                img_array[h//4:h//2, 0:w//4],
                img_array[h//4:h//2, w//4:w//2]
            ]
            
            noise_levels = []
            for region in regions:
                if len(region.shape) == 3:
                    gray = cv2.cvtColor(region, cv2.COLOR_RGB2GRAY)
                else:
                    gray = region
                noise = np.std(gray)
                noise_levels.append(noise)
            
            # Check for inconsistency
            avg_noise = np.mean(noise_levels)
            inconsistent = any(abs(noise - avg_noise) / avg_noise > 0.5 for noise in noise_levels)
            
            return {
                'noise_levels': noise_levels.tolist(),
                'average_noise': float(avg_noise),
                'inconsistent': inconsistent
            }
            
        except Exception as e:
            logger.error(f"Noise analysis failed: {str(e)}")
            return {'inconsistent': False}
    
    def _analyze_image_metadata(self, image: 'Image') -> Dict[str, Any]:
        """Analyze image metadata for suspicious indicators"""
        indicators = []
        suspicious = False
        
        try:
            # Check for missing metadata
            if not image.info:
                indicators.append("No EXIF metadata found")
                suspicious = True
            
            # Check for software editing signatures
            software = image.info.get('Software', '')
            if any(editor in software.lower() for editor in ['photoshop', 'gimp', 'paint.net']):
                indicators.append(f"Image edited with: {software}")
                suspicious = True
            
            # Check for unusual timestamps
            if 'DateTime' in image.info:
                datetime_str = image.info['DateTime']
                try:
                    from datetime import datetime, timezone
                    img_datetime = datetime.strptime(datetime_str, '%Y:%m:%d %H:%M:%S')
                    # Check if timestamp is in the future
                    if img_datetime > datetime.now():
                        indicators.append("Future timestamp detected")
                        suspicious = True
                except:
                    indicators.append("Invalid datetime format")
                    suspicious = True
            
            return {
                'indicators': indicators,
                'suspicious': suspicious,
                'metadata': dict(image.info)
            }
            
        except Exception as e:
            logger.error(f"Metadata analysis failed: {str(e)}")
            return {'indicators': [], 'suspicious': False}
    
    def _detect_clone_regions(self, image: 'Image') -> Dict[str, Any]:
        """Detect clone regions in image (simplified)"""
        try:
            import cv2
            import numpy as np
            
            # Convert to grayscale
            img_array = np.array(image.convert('L'))
            
            # Use template matching to find similar regions
            h, w = img_array.shape
            template_size = min(50, h//10, w//10)
            
            clones_detected = False
            # This is a simplified implementation
            # In practice, you'd use more sophisticated clone detection algorithms
            
            return {
                'clones_detected': clones_detected,
                'template_size': template_size
            }
            
        except Exception as e:
            logger.error(f"Clone detection failed: {str(e)}")
            return {'clones_detected': False}
    
    def _has_suspicious_pdf_metadata(self, metadata: Dict[str, Any]) -> bool:
        """Check for suspicious PDF metadata"""
        suspicious_indicators = []
        
        # Check for empty or generic metadata
        if not metadata or all(not value for value in metadata.values()):
            suspicious_indicators.append("Empty or generic metadata")
        
        # Check for unusual software
        software = metadata.get('/Producer', '')
        if any(suspicious in software.lower() for suspicious in ['unknown', 'fake', 'anonymous']):
            suspicious_indicators.append("Suspicious software in metadata")
        
        return len(suspicious_indicators) > 0
    
    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            import pytesseract
            import cv2
            
            # Preprocess image for better OCR
            image = cv2.imread(file_path)
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply threshold
            _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # OCR
            text = pytesseract.image_to_string(thresh, lang='eng')
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {str(e)}")
            return ""
    
    def _extract_entities_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract named entities from text (simplified)"""
        entities = []
        
        # Simple pattern matching for common entities
        import re
        
        # Email addresses
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        for email in set(emails):
            entities.append({'type': 'email', 'value': email, 'confidence': 0.9})
        
        # Phone numbers
        phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
        for phone in set(phones):
            entities.append({'type': 'phone', 'value': phone, 'confidence': 0.8})
        
        # URLs
        urls = re.findall(r'https?://[^\s<>"{}|\\^`]+', text)
        for url in set(urls):
            entities.append({'type': 'url', 'value': url, 'confidence': 0.9})
        
        return entities

    def _generate_evidence_id(self, file_path: str) -> str:
        import hashlib
        import uuid
        try:
             # Use timestamp + filename hash
             return f"ev_{hashlib.md5(f'{file_path}{datetime.now().timestamp()}'.encode()).hexdigest()[:12]}"
        except:
             return f"ev_{str(uuid.uuid4())[:12]}"

    def _detect_file_type(self, file_path: str) -> str:
        try:
            import magic
            return magic.from_file(file_path, mime=True)
        except ImportError:
            # Fallback based on extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext in ['.jpg', '.jpeg', '.png']: return 'image/jpeg'
            if ext == '.pdf': return 'application/pdf'
            if ext == '.txt': return 'text/plain'
            return 'application/octet-stream'
        except Exception:
            return 'application/octet-stream'

    def _identify_document_type(self, analysis: MultiModalAnalysis) -> str:
        text = analysis.extracted_text.lower()
        if 'bank statement' in text or 'account summary' in text:
            return 'bank_statement'
        if 'invoice' in text or 'bill to' in text:
            return 'invoice'
        return 'general_document'

    def _map_bank_statement_data(self, analysis: MultiModalAnalysis) -> Dict[str, Any]:
        return {}
    
    def _map_expense_data(self, analysis: MultiModalAnalysis) -> Dict[str, Any]:
        return {}

    def _calculate_quality_scores(self, analysis: MultiModalAnalysis):
        # Placeholder logic
        analysis.quality_score = 0.8
        analysis.relevance_score = 0.5
        analysis.admissibility_score = 0.9

    def _analyze_spreadsheet_file(self, file_path: str, analysis: MultiModalAnalysis):
        pass

# Singleton instance
multimodal_analyzer = MultiModalAnalyzer()