"""
Document Metadata Extraction Service

EXIF-like metadata extraction for fraud investigation documents.
Supports PDF, images (EXIF), and Office documents.
"""

import hashlib
import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Any

from pydantic import BaseModel

logger = logging.getLogger(__name__)


class DocumentHash(BaseModel):
    """File hash for chain of custody."""

    md5: str
    sha256: str


class CreationContext(BaseModel):
    """Creation metadata similar to EXIF."""

    date: str | None = None
    timezone: str | None = None
    software: str | None = None
    author: str | None = None
    device: str | None = None


class ModificationEvent(BaseModel):
    """Single modification event."""

    date: str
    action: str
    details: str | None = None


class ModificationHistory(BaseModel):
    """Document modification history."""

    last_date: str | None = None
    count: int = 0
    history: list[ModificationEvent] = []


class GeoLocation(BaseModel):
    """Geographic location if available."""

    lat: float | None = None
    lng: float | None = None
    accuracy: float | None = None
    source: str | None = None  # "GPS" | "IP" | "manual"


class PrintMetadata(BaseModel):
    """Print/scan metadata."""

    printer_name: str | None = None
    print_date: str | None = None
    copies: int | None = None


class PDFMetadata(BaseModel):
    """PDF-specific metadata."""

    producer: str | None = None
    version: str | None = None
    pages: int | None = None
    encrypted: bool = False
    permissions: list[str] = []


class CameraMetadata(BaseModel):
    """Camera EXIF data."""

    make: str | None = None
    model: str | None = None
    exposure: str | None = None
    iso: int | None = None


class ImageMetadata(BaseModel):
    """Image-specific EXIF metadata."""

    width: int | None = None
    height: int | None = None
    color_space: str | None = None
    dpi: int | None = None
    camera: CameraMetadata | None = None
    has_exif: bool = False
    file_format: str | None = None
    bits_per_pixel: int | None = None


class OfficeMetadata(BaseModel):
    """Office document metadata (DOCX, etc.)."""

    author: str | None = None
    created_date: datetime | None = None
    modified_date: datetime | None = None
    title: str | None = None
    subject: str | None = None
    keywords: str | None = None
    word_count: int = 0
    page_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0
    image_count: int = 0
    language: str = "en"
    revision_count: int = 0
    last_modified_by: str | None = None
    custom_properties: dict[str, Any] = {}


class ForensicFlags(BaseModel):
    """Forensic analysis flags."""

    tamper_likelihood: float = 0.0  # 0-100%
    anomalies: list[str] = []
    signature_valid: bool | None = None
    ocr_confidence: float | None = None


class DocumentMetadata(BaseModel):
    """Complete document metadata schema."""

    id: str
    filename: str
    filetype: str
    size: int
    hash: DocumentHash
    created: CreationContext
    modified: ModificationHistory
    location: GeoLocation | None = None
    print_info: PrintMetadata | None = None
    pdf: PDFMetadata | None = None
    image: ImageMetadata | None = None
    docx: OfficeMetadata | None = None
    forensic: ForensicFlags


class MetadataExtractionService:
    """
    Service for extracting EXIF-like metadata from documents.
    """

    def __init__(self):
        self.supported_types = {
            "application/pdf": self._extract_pdf_metadata,
            "image/jpeg": self._extract_image_metadata,
            "image/png": self._extract_image_metadata,
            "image/tiff": self._extract_image_metadata,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx_metadata,
        }

    def calculate_hash(self, file_path: Path) -> DocumentHash:
        """Calculate MD5 and SHA-256 hashes for chain of custody."""
        md5_hash = hashlib.md5()
        sha256_hash = hashlib.sha256()

        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)
                sha256_hash.update(chunk)

        return DocumentHash(md5=md5_hash.hexdigest(), sha256=sha256_hash.hexdigest())

    def extract_metadata(self, file_path: Path, ocr_result: dict = None) -> DocumentMetadata:
        """
        Extract all available metadata from a document.

        Args:
            file_path: Path to the document file
            ocr_result: Optional OCR processing result from evidence service

        Returns:
            DocumentMetadata with all extracted information
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Basic file info
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        file_hash = self.calculate_hash(file_path)

        # Extract OCR confidence from ocr_result if available
        ocr_confidence = None
        if ocr_result and 'metadata' in ocr_result:
            ocr_confidence = ocr_result['metadata'].get('ocr_confidence')

        # Base metadata
        metadata = DocumentMetadata(
            id=file_hash.sha256[:16],
            filename=file_path.name,
            filetype=mime_type or "application/octet-stream",
            size=stat.st_size,
            hash=file_hash,
            created=CreationContext(
                date=datetime.fromtimestamp(stat.st_ctime).isoformat(),
            ),
            modified=ModificationHistory(
                last_date=datetime.fromtimestamp(stat.st_mtime).isoformat(), count=1
            ),
            forensic=ForensicFlags(
                ocr_confidence=ocr_confidence
            ),
        )

        # Type-specific extraction
        if mime_type in self.supported_types:
            type_metadata = self.supported_types[mime_type](file_path)
            metadata = self._merge_metadata(metadata, type_metadata)

        return metadata

    def _extract_pdf_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract PDF-specific metadata using pypdf (pdf-lib equivalent)."""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                return {
                    "pdf": PDFMetadata(
                        producer="PDF library not available",
                        version="unknown",
                        pages=0,
                        encrypted=False,
                        permissions=[]
                    )
                }

        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)

                # Extract basic metadata
                metadata = reader.metadata
                num_pages = len(reader.pages)

                # Extract producer info
                producer = getattr(metadata, 'producer', None) or getattr(metadata, '/Producer', None)
                if isinstance(producer, bytes):
                    producer = producer.decode('utf-8', errors='ignore')

                # Extract version info
                version = "1.4"  # Default
                if hasattr(reader, 'pdf_header'):
                    header = reader.pdf_header
                    if b'PDF-1.' in header:
                        version_match = header.split(b'PDF-1.')[1][:1]
                        if version_match.isdigit():
                            version = f"1.{version_match.decode()}"

                # Check encryption
                encrypted = reader.is_encrypted

                # Extract permissions (if not encrypted)
                permissions = []
                if not encrypted:
                    try:
                        # Check if we can extract text (implies read permission)
                        page = reader.pages[0]
                        if hasattr(page, 'extract_text'):
                            permissions.append("read")
                    except Exception:
                        pass

                    # Check if we can access page objects
                    try:
                        _ = reader.pages[0]
                        permissions.append("print")  # Basic print permission
                    except Exception:
                        pass

                # Extract additional metadata
                creation_date = getattr(metadata, 'creation_date', None)
                modification_date = getattr(metadata, 'modification_date', None)
                author = getattr(metadata, 'author', None)
                subject = getattr(metadata, 'subject', None)
                title = getattr(metadata, 'title', None)

                # Convert bytes to strings if needed
                for attr in ['author', 'subject', 'title']:
                    value = locals().get(attr)
                    if isinstance(value, bytes):
                        locals()[attr] = value.decode('utf-8', errors='ignore')

                return {
                    "pdf": PDFMetadata(
                        producer=producer or "Unknown",
                        version=version,
                        pages=num_pages,
                        encrypted=encrypted,
                        permissions=permissions,
                        creation_date=creation_date,
                        modification_date=modification_date,
                        author=author,
                        subject=subject,
                        title=title
                    )
                }

        except Exception as e:
            return {
                "pdf": PDFMetadata(
                    producer=f"Error extracting PDF metadata: {e!s}",
                    version="unknown",
                    pages=0,
                    encrypted=False,
                    permissions=[]
                )
            }

    def _extract_image_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract image EXIF metadata using Pillow."""
        try:
            from PIL import ExifTags, Image
        except ImportError:
            return {
                "exif": EXIFMetadata(
                    camera_make="Pillow not available",
                    camera_model="PIL library required",
                    datetime_original=None,
                    gps_latitude=None,
                    gps_longitude=None,
                ),
                "image": ImageMetadata(
                    width=0,
                    height=0,
                    color_space="unknown",
                    has_exif=False,
                )
            }

        try:
            with Image.open(file_path) as img:
                # Get basic image info
                width, height = img.size
                color_space = "RGB" if img.mode == "RGB" else img.mode
                has_exif = hasattr(img, '_getexif') and img._getexif() is not None

                # Extract EXIF data
                exif_data = {}
                if has_exif:
                    exif_dict = img._getexif()
                    if exif_dict:
                        for tag, value in exif_dict.items():
                            tag_name = ExifTags.TAGS.get(tag, tag)
                            exif_data[tag_name] = value

                # Extract camera info
                camera_make = exif_data.get('Make', 'Unknown')
                camera_model = exif_data.get('Model', 'Unknown')

                # Extract datetime
                datetime_original = exif_data.get('DateTimeOriginal')
                if isinstance(datetime_original, str):
                    try:
                        # Convert EXIF datetime format to datetime object
                        from datetime import datetime
                        datetime_original = datetime.strptime(datetime_original, '%Y:%m:%d %H:%M:%S')
                    except Exception:
                        datetime_original = None

                # Extract GPS data
                gps_latitude = None
                gps_longitude = None

                if 'GPSInfo' in exif_data:
                    gps_info = exif_data['GPSInfo']

                    # GPS latitude
                    if 2 in gps_info and 4 in gps_info:
                        lat_deg = gps_info[2][0] / gps_info[2][1]
                        lat_min = gps_info[2][1] / gps_info[2][2]
                        lat_sec = gps_info[2][2] / gps_info[2][3]
                        gps_latitude = lat_deg + (lat_min / 60) + (lat_sec / 3600)

                        if gps_info[1] == 'S':
                            gps_latitude = -gps_latitude

                    # GPS longitude
                    if 4 in gps_info and 6 in gps_info:
                        lon_deg = gps_info[4][0] / gps_info[4][1]
                        lon_min = gps_info[4][1] / gps_info[4][2]
                        lon_sec = gps_info[4][2] / gps_info[4][3]
                        gps_longitude = lon_deg + (lon_min / 60) + (lon_sec / 3600)

                        if gps_info[3] == 'W':
                            gps_longitude = -gps_longitude

                # Additional EXIF fields
                iso_speed = exif_data.get('ISOSpeedRatings')
                focal_length = exif_data.get('FocalLength')
                aperture = exif_data.get('FNumber')
                exposure_time = exif_data.get('ExposureTime')
                flash = exif_data.get('Flash')

                return {
                    "exif": EXIFMetadata(
                        camera_make=camera_make,
                        camera_model=camera_model,
                        datetime_original=datetime_original,
                        gps_latitude=gps_latitude,
                        gps_longitude=gps_longitude,
                        iso_speed=iso_speed,
                        focal_length=focal_length,
                        aperture=aperture,
                        exposure_time=exposure_time,
                        flash_used=flash == 1 if flash is not None else None
                    ),
                    "image": ImageMetadata(
                        width=width,
                        height=height,
                        color_space=color_space,
                        has_exif=has_exif,
                        file_format=img.format,
                        bits_per_pixel=getattr(img, 'bits', None)
                    )
                }

        except Exception as e:
            return {
                "exif": EXIFMetadata(
                    camera_make=f"Error extracting EXIF: {e!s}",
                    camera_model="Error",
                    datetime_original=None,
                    gps_latitude=None,
                    gps_longitude=None,
                ),
                "image": ImageMetadata(
                    width=0,
                    height=0,
                    color_space="unknown",
                    has_exif=False,
                )
            }

    def _extract_docx_metadata(self, file_path: Path) -> dict[str, Any]:
        """Extract DOCX metadata using python-docx."""
        try:
            import zipfile
            from datetime import datetime

            from docx import Document
        except ImportError:
            return {
                "docx": OfficeMetadata(
                    author="python-docx not available",
                    created_date=None,
                    modified_date=None,
                    word_count=0,
                    page_count=0,
                )
            }

        try:
            # Load the document
            doc = Document(file_path)

            # Extract core properties
            core_props = doc.core_properties

            author = getattr(core_props, 'author', None) or 'Unknown'
            created_date = getattr(core_props, 'created', None)
            modified_date = getattr(core_props, 'modified', None)
            title = getattr(core_props, 'title', None)
            subject = getattr(core_props, 'subject', None)
            keywords = getattr(core_props, 'keywords', None)

            # Count words and paragraphs
            word_count = 0
            paragraph_count = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_count += 1
                    word_count += len(paragraph.text.split())

            # Estimate page count (rough approximation)
            # Average 300 words per page for typical documents
            page_count = max(1, word_count // 300)

            # Extract additional metadata from document structure
            table_count = len(doc.tables)
            image_count = 0

            # Count images by checking relationships
            try:
                rels = doc.part.rels
                for rel in rels.values():
                    if hasattr(rel, 'target_ref') and rel.target_ref:
                        if any(ext in rel.target_ref.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']):
                            image_count += 1
            except Exception:
                pass

            # Extract custom properties if available
            custom_props = {}
            try:
                # Some versions of python-docx support custom properties
                if hasattr(doc, 'custom_properties'):
                    for prop in doc.custom_properties:
                        custom_props[prop.name] = str(prop.value)
            except Exception:
                pass

            # Determine document language (basic heuristic)
            language = "en"  # Default
            # Could be enhanced with language detection libraries

            # Extract revision history (basic)
            revision_count = 0
            last_modified_by = getattr(core_props, 'last_modified_by', None)

            return {
                "docx": OfficeMetadata(
                    author=author,
                    created_date=created_date,
                    modified_date=modified_date,
                    title=title,
                    subject=subject,
                    keywords=keywords,
                    word_count=word_count,
                    page_count=page_count,
                    paragraph_count=paragraph_count,
                    table_count=table_count,
                    image_count=image_count,
                    language=language,
                    revision_count=revision_count,
                    last_modified_by=last_modified_by,
                    custom_properties=custom_props
                )
            }

        except Exception as e:
            return {
                "docx": OfficeMetadata(
                    author=f"Error extracting DOCX metadata: {e!s}",
                    created_date=None,
                    modified_date=None,
                    word_count=0,
                    page_count=0,
                )
            }

    def _merge_metadata(
        self, base: DocumentMetadata, additional: dict[str, Any]
    ) -> DocumentMetadata:
        """Merge additional metadata into base."""
        data = base.model_dump()
        for key, value in additional.items():
            if value is not None:
                if isinstance(value, BaseModel):
                    data[key] = value.model_dump()
                else:
                    data[key] = value
        return DocumentMetadata(**data)

    def compare_documents(
        self, doc_a: DocumentMetadata, doc_b: DocumentMetadata
    ) -> dict[str, Any]:
        """
        Compare two documents and detect discrepancies.

        Returns dict with:
        - matches: List of matching fields
        - discrepancies: List of different fields with details
        - tamper_indicators: List of potential tampering signs
        """
        discrepancies = []
        tamper_indicators = []

        # Compare hashes
        if doc_a.hash.sha256 != doc_b.hash.sha256:
            discrepancies.append(
                {
                    "field": "content_hash",
                    "doc_a": doc_a.hash.sha256[:16] + "...",
                    "doc_b": doc_b.hash.sha256[:16] + "...",
                    "severity": "high",
                }
            )
            tamper_indicators.append("Content modified between versions")

        # Compare authors
        if doc_a.created.author != doc_b.created.author:
            discrepancies.append(
                {
                    "field": "author",
                    "doc_a": doc_a.created.author,
                    "doc_b": doc_b.created.author,
                    "severity": "medium",
                }
            )
            tamper_indicators.append("Author name changed")

        # Compare software
        if doc_a.created.software != doc_b.created.software:
            discrepancies.append(
                {
                    "field": "software",
                    "doc_a": doc_a.created.software,
                    "doc_b": doc_b.created.software,
                    "severity": "medium",
                }
            )
            tamper_indicators.append('Different software used for "same" document')

        # Check modification timing
        if doc_a.modified.last_date and doc_b.modified.last_date:
            a_date = datetime.fromisoformat(doc_a.modified.last_date)
            b_date = datetime.fromisoformat(doc_b.modified.last_date)
            if (b_date - a_date).days > 1:
                tamper_indicators.append(
                    f"Modified {(b_date - a_date).days} days after original"
                )

        return {
            "hash_match": doc_a.hash.sha256 == doc_b.hash.sha256,
            "discrepancies": discrepancies,
            "tamper_indicators": tamper_indicators,
            "risk_score": len(tamper_indicators) * 25,  # 0-100
        }

    def detect_tampering(self, metadata: DocumentMetadata) -> ForensicFlags:
        """
        Analyze metadata for signs of tampering.

        Returns updated ForensicFlags with analysis results.
        """
        anomalies = []
        tamper_likelihood = 0.0

        # Check for metadata inconsistencies
        if metadata.created.date and metadata.modified.last_date:
            created = datetime.fromisoformat(metadata.created.date)
            modified = datetime.fromisoformat(metadata.modified.last_date)

            if modified < created:
                anomalies.append("modification_before_creation")
                tamper_likelihood += 30

        # Check for suspicious software
        if metadata.created.software:
            suspicious_editors = ["photoshop", "gimp", "acrobat pro"]
            if any(s in metadata.created.software.lower() for s in suspicious_editors):
                anomalies.append("editing_software_detected")
                tamper_likelihood += 15

        # Check for missing expected metadata
        if not metadata.created.author:
            anomalies.append("missing_author")
            tamper_likelihood += 10

        return ForensicFlags(
            tamper_likelihood=min(tamper_likelihood, 100),
            anomalies=anomalies,
            signature_valid=None,  # Requires digital signature check
            ocr_confidence=0.0,  # Will be updated by OCR analysis
        )


# Create singleton instance
metadata_service = MetadataExtractionService()
